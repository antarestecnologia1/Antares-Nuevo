# -*- coding: utf-8 -*-
"""Genera Antares_Arquitectura_UI_UX.docx — arquitectura AS-IS validada contra el repositorio."""
from __future__ import annotations

import argparse
import subprocess
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

_DOCS = Path(__file__).resolve().parent
_ROOT = Path(__file__).resolve().parents[1]
if str(_DOCS) not in sys.path:
    sys.path.insert(0, str(_DOCS))
from antares_diagrams_matplotlib import generate_all_diagrams

# Paleta corporativa Antares
C_BRAND = RGBColor(0x37, 0x7C, 0xC0)
C_BRAND_DARK = RGBColor(0x1E, 0x4A, 0x73)
C_TEXT = RGBColor(0x0B, 0x21, 0x38)
C_MUTED = RGBColor(0x3A, 0x5A, 0x78)
C_SUCCESS = RGBColor(0x1B, 0x8E, 0x5F)
C_WARN = RGBColor(0xF5, 0x9F, 0x00)
C_FILL_HDR = "377CC0"
C_FILL_ALT = "E8F4FC"
FECHA_INFORME = date(2026, 5, 16)
VERSION = "3.0"


def _logo_path() -> Path | None:
    for p in (
        _ROOT / "imagenes empresa" / "Logo.png",
        _ROOT / "imagenes%20empresa" / "Logo.png",
    ):
        if p.is_file():
            return p
    return None


def _mes_es(d: date) -> str:
    meses = (
        "enero", "febrero", "marzo", "abril", "mayo", "junio",
        "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
    )
    return f"{d.day} de {meses[d.month - 1]} de {d.year}"


def _set_cell_shading(cell, fill_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _cell_write(cell, text: str, *, bold=False, size=10, color: RGBColor | None = None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Calibri"
    if color:
        r.font.color.rgb = color


def add_h(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = C_BRAND_DARK if level == 1 else C_BRAND
    return h


def add_p(doc: Document, text: str, *, bold=False, italic=False, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(11)
    r.font.name = "Calibri"
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_callout(doc: Document, title: str, body: str):
    """Bloque destacado estilo «nota ejecutiva»."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    _set_cell_shading(cell, C_FILL_ALT)
    _cell_write(cell, title, bold=True, size=11, color=C_BRAND_DARK)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(4)
    r2 = p2.add_run(body)
    r2.font.size = Pt(10.5)
    r2.font.name = "Calibri"
    r2.font.color.rgb = C_TEXT
    doc.add_paragraph()


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    *,
    col_widths: list[float] | None = None,
    zebra: bool = True,
):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        _set_cell_shading(hdr.cells[i], C_FILL_HDR)
        _cell_write(hdr.cells[i], h, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            if zebra and ri % 2 == 1:
                _set_cell_shading(cells[ci], C_FILL_ALT)
            _cell_write(cells[ci], val, size=9.5)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def add_figure(doc: Document, image_path: Path, caption_es: str, width_in: float = 6.45):
    if not image_path.is_file():
        add_p(doc, f"[Diagrama no generado — {image_path.name}]", bold=True)
        return
    pg = doc.add_paragraph()
    pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pg.paragraph_format.space_before = Pt(8)
    pg.add_run().add_picture(str(image_path), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(14)
    r = cap.add_run(caption_es)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = C_MUTED


def add_mono(doc: Document, text: str, size_pt: int = 8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(size_pt)
    r.font.color.rgb = C_MUTED


def add_cover_page(doc: Document) -> Path | None:
    logo = _logo_path()
    for _ in range(2):
        doc.add_paragraph()
    if logo:
        pg = doc.add_paragraph()
        pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pg.add_run().add_picture(str(logo), width=Inches(4.1))

    band = doc.add_table(rows=1, cols=1)
    band.style = "Table Grid"
    bc = band.rows[0].cells[0]
    _set_cell_shading(bc, C_FILL_HDR)
    bc.text = ""
    for line, pt, bold, color in [
        ("TRANSPORTES ANTARES S.A.S.", 22, True, RGBColor(0xFF, 0xFF, 0xFF)),
        ("Plataforma Web Integral — Arquitectura técnica y experiencia de usuario", 13, True, RGBColor(0xE8, 0xF4, 0xFC)),
        ("Estado AS-IS validado en código fuente · Mayo 2026", 11, False, RGBColor(0xCC, 0xE5, 0xF8)),
    ]:
        p = bc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        r.bold = bold
        r.font.size = Pt(pt)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta.add_run(
        f"Versión {VERSION}  ·  {_mes_es(FECHA_INFORME)}  ·  Clasificación: Confidencial"
    )
    rm.font.size = Pt(10)
    rm.italic = True
    rm.font.color.rgb = C_MUTED

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = tag.add_run("Entregable institucional — Comité técnico · Dirección · Cliente")
    rt.bold = True
    rt.font.size = Pt(10)
    rt.font.color.rgb = RGBColor(0xD6, 0x28, 0x28)

    doc.add_page_break()
    return logo


def add_toc_page(doc: Document):
    add_h(doc, "Índice", 1)
    entries = [
        ("Resumen ejecutivo y métricas del sistema", "3"),
        ("Hoja de control documental", "4"),
        ("1. Propósito, audiencia y actores", "5"),
        ("2. Mapa funcional por módulo contractual", "6"),
        ("3. Portal operativo — vistas y roles", "8"),
        ("4. Marcos de referencia y diagramas C4-lite", "10"),
        ("5. Stack tecnológico y despliegue", "16"),
        ("6. Sistema de diseño UI/UX", "17"),
        ("7. Sincronización datos y API REST", "19"),
        ("8. Modelo de datos PostgreSQL", "21"),
        ("9. Atributos de calidad y pruebas", "22"),
        ("10. Roadmap, limitaciones y glosario", "23"),
    ]
    add_table(doc, ["Sección", "Pág. ref."], entries, zebra=False)
    add_p(
        doc,
        "Nota: al abrir en Microsoft Word puede insertar tabla de contenido automática "
        "(Referencias → Tabla de contenido) usando los estilos Título 1–3 de este documento.",
        italic=True,
        space_after=12,
    )
    doc.add_page_break()


def _count_sql_migrations() -> int:
    pg = _ROOT / "BD" / "postgres"
    if not pg.is_dir():
        return 0
    return len([p for p in pg.glob("*.sql") if p.name[0:2].isdigit()])


def _portal_views() -> list[list[str]]:
    """Vistas registradas en modules/portal/architecture.js (AS-IS)."""
    return [
        ["dashboard", "Dashboard", "Cualquier rol autenticado", "KPIs operativos y accesos rápidos"],
        ["requests", "Solicitudes", "Cliente · Admin", "Crear y listar solicitudes de transporte"],
        ["transport-trips", "Transporte · Viajes", "Admin", "Viajes activos y asignación"],
        ["transport-vehicles", "Transporte · Camiones", "Admin", "Flota, disponibilidad, satélite"],
        ["transport-drivers", "Transporte · Conductores", "Admin", "Licencias, fotos, disponibilidad"],
        ["transport-calendar", "Transporte · Calendario", "Admin", "Agenda y solapes horarios"],
        ["history", "Historial y reportes", "Admin", "Filtros por fecha, cliente, estado"],
        ["reports", "Centro de reportería", "Admin · RRHH", "Métricas operativas consolidadas"],
        ["payroll", "Gestión humana / Nómina", "RRHH", "Empleados, liquidación, desprendibles"],
        ["hiring", "Contratación", "RRHH", "Vacantes, candidatos, entrevistas, contratos DOCX"],
        ["labor-compliance", "Cumplimiento laboral y SST", "RRHH", "Exámenes ocupacionales, intravehicular"],
        ["admin-users", "Usuarios y permisos", "Admin", "Altas, roles, sesiones, credenciales"],
        ["authorizations", "Centro de aprobaciones", "Admin", "Ventana 10 min · auto-aprobación"],
        ["contact-leads", "Contacto web B2B", "Admin · permiso contact_b2b", "Prospectos formulario público"],
        ["profile", "Mi perfil", "Todos", "Datos de cuenta y preferencias"],
        ["notifications", "Notificaciones", "Todos", "Bandeja, sonido, preferencias push/correo"],
    ]


def _module_contract_matrix() -> list[list[str]]:
    return [
        [
            "Módulo 1 — Sitio corporativo público",
            "Operativo",
            "100 %",
            "Hero, about, servicios, flota, cobertura, noticias, empleos, contacto B2B, testimonios",
        ],
        [
            "Módulo 2 — Portal clientes empresariales",
            "Operativo",
            "95 %",
            "JWT, registro con aprobación admin, solicitudes, estados, notificaciones",
        ],
        [
            "Módulo 3 — Panel administración",
            "Operativo",
            "90 %",
            "Flota, conductores, viajes, calendario, historial, reportes, aprobaciones 10 min",
        ],
        [
            "Módulo 4 — Nómina RRHH",
            "Operativo",
            "85 %",
            "Liquidación CO, períodos extendidos, prima/terminación, PDF, exportación",
        ],
        [
            "Módulo 5 — Contratación RRHH",
            "Operativo",
            "80 %",
            "Vacantes, pipeline selección, CV descarga API, contratos desde plantilla",
        ],
        [
            "Transversal — API + datos + seguridad",
            "Operativo",
            "90 %",
            "NestJS 10, PostgreSQL 36+ migraciones, R2, Turnstile, RLS opcional Supabase",
        ],
    ]


def export_pdf_via_word(docx_path: Path, pdf_path: Path) -> tuple[bool, str]:
    wd_pdf = 17
    docx_abs = str(docx_path.resolve())
    pdf_abs = str(pdf_path.resolve())
    if pdf_path.is_file():
        pdf_path.unlink()
    trace = ""
    try:
        import win32com.client as win32

        app = win32.DispatchEx("Word.Application")
        app.Visible = False
        app.DisplayAlerts = 0
        d = None
        try:
            d = app.Documents.Open(docx_abs, ReadOnly=True)
            d.SaveAs(pdf_abs, FileFormat=wd_pdf)
        finally:
            if d:
                d.Close(SaveChanges=0)
            app.Quit(SaveChanges=0)
        if pdf_path.is_file():
            return True, "PDF vía Microsoft Word."
    except Exception as e:
        trace = str(e)
    try:
        from docx2pdf import convert

        convert(docx_abs, pdf_abs)
        if pdf_path.is_file():
            return True, "PDF vía docx2pdf."
    except Exception as e:
        trace = f"{trace} · {e}"
    return False, trace


def export_pdf_via_libreoffice(soffice: str | None, docx_path: Path, pdf_out_dir: Path) -> tuple[bool, str]:
    candidates = []
    if soffice:
        candidates.append(Path(soffice))
    candidates.extend([
        Path(r"C:\Program Files\LibreOffice\program\soffice.exe"),
        Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"),
    ])
    exe = next((p for p in candidates if p.is_file()), None)
    if not exe:
        return False, ""
    pdf_out_dir.mkdir(parents=True, exist_ok=True)
    res = subprocess.run(
        [str(exe), "--headless", "--norestore", "--nologo", "--nolockcheck",
         "--convert-to", "pdf", "--outdir", str(pdf_out_dir.resolve()), str(docx_path.resolve())],
        capture_output=True, text=True, check=False, timeout=180,
    )
    target = pdf_out_dir / docx_path.with_suffix(".pdf").name
    if target.is_file():
        return True, f"PDF vía LibreOffice ({exe.name})."
    return False, (res.stderr or "")[:400]


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--no-pdf", action="store_true")
    ap.add_argument("--libreoffice-first", action="store_true")
    ap.add_argument("--soffice", default="", metavar="RUTA")
    args = ap.parse_args()

    fig_dir = _DOCS / "_diagramas_antares_build"
    figs = generate_all_diagrams(fig_dir)
    n_sql = _count_sql_migrations()
    out = _DOCS / "Antares_Arquitectura_UI_UX.docx"

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    for level in range(1, 4):
        st = doc.styles[f"Heading {level}"]
        st.font.name = "Calibri"
        st.font.color.rgb = C_BRAND_DARK if level == 1 else C_BRAND

    logo = add_cover_page(doc)
    add_toc_page(doc)

    # —— Resumen ejecutivo ——
    add_h(doc, "Resumen ejecutivo", 1)
    add_callout(
        doc,
        "Hallazgo principal (validación " + _mes_es(FECHA_INFORME) + ")",
        "La plataforma Antares opera como solución full-stack integrada: sitio público B2B y portal "
        "autenticado multi-rol sobre un único SPA (`index.html` + `app.js` + `styles.css`), con "
        "persistencia autoritativa en PostgreSQL y API NestJS 10. Los cinco módulos del contrato "
        "tienen implementación funcional verificable; el trabajo restante se concentra en endurecimiento "
        "QA de producción, documentación OpenAPI formal y pulido de entregables finales.",
    )
    add_table(
        doc,
        ["Indicador", "Valor AS-IS"],
        [
            ["Migraciones SQL versionadas (`BD/postgres/`)", f"{n_sql} scripts numerados"],
            ["Vistas de portal registradas", f"{len(_portal_views())} rutas en PortalArchitecture"],
            ["Módulos NestJS (`AppModule`)", "Auth · Portal · Payroll · Mail · Files · Uploads · B2B"],
            ["Roles de negocio", "Admin · Cliente · RRHH · Administración · Auxiliar · Líder admin."],
            ["Cliente canónico en producción", "SPA raíz (`index.html`) — `apps/web` es workspace auxiliar"],
            ["Pruebas automatizadas portal", "`npm run test:portal-qa` · `npm run smoke:api`"],
        ],
    )
    add_bullets(
        doc,
        [
            "Arquitectura modular en el navegador: capas portal (`modules/portal/*`), dominio (`DomainRegistry`) y sincronización (`PortalDataLayer`, `portal-sync`).",
            "API `/api/portal/bootstrap` hidrata el estado tras login; escrituras incrementales vía `POST /api/portal/sync-key`.",
            "Seguridad: JWT + refresh rotativo, bcrypt, Cloudflare Turnstile, CSP, throttling 80 req/min, sesión BD en `America/Bogota`.",
            "Almacenamiento de documentos: UploadsModule con R2/S3 presigned; correo transaccional Resend.",
        ],
    )

    doc.add_page_break()
    add_h(doc, "Hoja de control documental", 1)
    add_table(
        doc,
        ["Campo", "Valor"],
        [
            ["Título", "Arquitectura técnica y UI/UX — Plataforma Antares (AS-IS)"],
            ["Versión", f"{VERSION} — validada contra repositorio «Antares Nuevo»"],
            ["Fecha de corte", _mes_es(FECHA_INFORME)],
            ["Clasificación", "Confidencial — Transportes Antares S.A.S. y equipo autorizado"],
            ["Generación", str(_DOCS / "generar_arquitectura_uiux.py")],
            ["Diagramas", f"Matplotlib 220 dpi → {fig_dir.name}/"],
            ["Logo portada", str(logo.resolve()) if logo else "No encontrado — revisar `imagenes empresa/Logo.png`"],
        ],
    )

    doc.add_page_break()
    add_h(doc, "1. Propósito, audiencia y actores", 1)
    add_p(
        doc,
        "Documento de arquitectura de solución en estado AS-IS: describe lo implementado y observable "
        "en el código, no un diseño objetivo futuro. Complementa el contrato y el Anexo A sin sustituir "
        "actas de recibo ni políticas corporativas de seguridad.",
    )
    add_table(
        doc,
        ["Actor", "Preocupación", "Componentes que atiende"],
        [
            ["Visitante / prospecto B2B", "Confianza de marca, contacto comercial", "Sitio público · formulario prospectos"],
            ["Usuario empresa (cliente)", "Solicitar y rastrear transporte de carga", "Portal solicitudes · notificaciones"],
            ["Operaciones / Admin", "Asignar flota, aprobar en ventana legal", "Transporte · autorizaciones · reportes"],
            ["RRHH", "Nómina legal Colombia · contratación", "Payroll · hiring · cumplimiento SST"],
            ["TI / Seguridad", "Superficie de ataque, secretos, auditoría", "API NestJS · RLS · headers · rate limit"],
        ],
    )

    doc.add_page_break()
    add_h(doc, "2. Mapa funcional por módulo contractual", 1)
    add_p(doc, "Matriz de madurez alineada al Anexo A del contrato de desarrollo (cinco módulos + transversal).")
    add_table(
        doc,
        ["Módulo", "Estado", "Avance ref.", "Capacidades verificadas en código"],
        _module_contract_matrix(),
        col_widths=[1.55, 0.75, 0.65, 3.55],
    )

    add_h(doc, "2.1 Sitio público — secciones (`index.html`)", 2)
    add_table(
        doc,
        ["Anchor", "Sección", "Contenido"],
        [
            ["#hero", "Hero", "Propuesta de valor · CTA contacto / portal"],
            ["#about", "Quiénes somos", "Misión, visión, valores (CMS vía admin)"],
            ["#trusted", "Confianza", "Logos / partners"],
            ["#testimonials", "Testimonios", "Carrusel reseñas clientes"],
            ["#services", "Servicios", "Transporte general y especializado"],
            ["#fleet", "Flota", "Turbo, camión, tractocamión, refrigeración"],
            ["#process", "Proceso", "Pasos operativos de contratación"],
            ["#coverage", "Cobertura", "Rutas y territorios"],
            ["#news", "Noticias", "Novedades corporativas"],
            ["#careers", "Empleos", "Enlace a vacantes / portal RRHH"],
            ["#contact", "Contacto B2B", "Formulario → API `b2b-prospect`"],
        ],
    )

    doc.add_page_break()
    add_h(doc, "3. Portal operativo — vistas y roles", 1)
    add_table(
        doc,
        ["ID vista", "Título en UI", "Acceso", "Función"],
        _portal_views(),
        col_widths=[1.15, 1.35, 1.05, 2.95],
    )
    add_p(
        doc,
        "Patrones UX transversales del portal: shell modular con `fleet-hero-strip`, barras de comando "
        "`.ops-command-bar`, tablas responsivas `.table-wrap`, estados vacíos `.empty-state`, "
        "tema claro/oscuro `data-theme` y navegación lateral con badges de notificación.",
    )

    doc.add_page_break()
    add_h(doc, "4. Marcos de referencia y diagramas", 1)
    add_p(
        doc,
        "Vistas inspiradas en ISO/IEC/IEEE 42010 y notación C4 (niveles 1–3 simplificados). "
        "Figuras generadas automáticamente para auditoría reproducible.",
    )

    add_h(doc, "4.1 Vista contextual", 2)
    add_figure(doc, figs["fig_a_contexto_sistema"], "Figura A — Actores, frontera del sistema y sistemas externos.")
    add_h(doc, "4.2 Vista de contenedores", 2)
    add_figure(doc, figs["fig_b_contenedores"], "Figura B — SPA, CDN, API NestJS, PostgreSQL y object storage.")
    doc.add_page_break()
    add_h(doc, "4.3 Vista de despliegue", 2)
    add_figure(doc, figs["fig_c_despliegue_nube"], "Figura C — Patrón Cloudflare + API + Postgres administrado (Supabase/Render).")
    add_h(doc, "4.4 Componentes backend", 2)
    add_figure(doc, figs["fig_d_componentes_api"], "Figura D — Módulos NestJS 10 desplegados bajo prefijo `/api`.")
    add_h(doc, "4.5 Contextos de dominio / datos", 2)
    add_figure(doc, figs["fig_e_contextos_datos"], "Figura E — Bounded contexts alineados al DDL `BD/postgres/`.")
    doc.add_page_break()
    add_h(doc, "4.6 Seguridad en profundidad", 2)
    add_figure(doc, figs["fig_f_capas_seguridad"], "Figura F — Controles desde perímetro CDN hasta integridad en BD.")

    doc.add_page_break()
    add_h(doc, "5. Stack tecnológico y despliegue", 1)
    add_table(
        doc,
        ["Capa", "Tecnología AS-IS", "Ubicación / notas"],
        [
            ["Frontend principal", "HTML5 · CSS3 · JavaScript ES modules", "`index.html` · `app.js` · `styles.css` · `modules/`"],
            ["Frontend auxiliar", "Next.js (workspace `apps/web`)", "Módulos espejo parciales — no canónico en prod."],
            ["Backend", "Node.js 20+ · NestJS 10 · TypeScript", "`apps/api` — build `tsc`"],
            ["Base de datos", "PostgreSQL 15+ · `pg` pool", "Scripts `BD/postgres/` · TZ sesión Bogotá"],
            ["Auth", "Passport JWT · refresh hash en `usuarios`", "`apps/api/src/auth`"],
            ["Archivos", "AWS SDK S3 · Cloudflare R2", "`apps/api/src/uploads/r2.service.ts`"],
            ["Correo", "Resend", "`MailModule`"],
            ["Edge / DNS", "Cloudflare Pages + Tunnel (documentado)", "`.cfignore` · runbooks en `docs/`"],
            ["Opcional BD host", "Supabase (Auth, Storage, RLS)", "Scripts `09_rls_tablas.sql`, `10_rls_storage`"],
        ],
    )

    add_h(doc, "6. Sistema de diseño UI/UX institucional", 1)
    add_p(doc, "Tokens de marca centralizados en `:root` de `styles.css` — coherencia entre sitio público y portal.")
    add_table(
        doc,
        ["Token / rol", "Valor", "Uso"],
        [
            ["--brand-blue-deep", "#377CC0", "Primario · botones · acentos módulo"],
            ["--brand-blue-mid", "#83BEE9", "Gradientes · hover · acento secundario"],
            ["--gradient-hero", "135° deep → mid", "Hero sitio público"],
            ["--font-body", "Montserrat", "Texto general y UI densa"],
            ["--font-display", "Poppins", "Títulos y jerarquía visual"],
            ["--font-tertiary", "Roboto / Lato", "Datos tabulares y microcopy"],
            ["--radius", "16px", "Tarjetas glass / módulos portal"],
            ["Modo oscuro", "`body[data-theme=\"dark\"]`", "Portal y sitio — contraste AA orientativo"],
        ],
    )
    add_h(doc, "6.1 Arquitectura de información — sitio público", 2)
    add_figure(
        doc,
        figs["fig_g_ia_sitio_publico"],
        "Figura G — Secuencia de anclas y scroll del menú principal.",
        width_in=6.95,
    )
    add_bullets(
        doc,
        [
            "Navegación sticky con blur (`--public-nav-bg`) y menú móvil con panel deslizable.",
            "Componentes reutilizables: `.glass-card`, `.hero`, `.section-dark`, grids de flota.",
            "Formulario B2B con validación cliente (`modules/core/validation.js`) y anti-bot Turnstile.",
            "Portal: overlay de autenticación, sidebar por rol, shells `.module-shell` por vista.",
        ],
    )

    doc.add_page_break()
    add_h(doc, "7. Sincronización de datos y API REST", 1)
    add_table(
        doc,
        ["Colección lógica (KEYS)", "Tabla(s) PostgreSQL", "Notas"],
        [
            ["users", "usuarios", "Roles, permisos_usuario, refresh_token_hash"],
            ["companies", "empresas", "Logo URL, contacto, tipo relación, activo"],
            ["requests", "solicitudes_transporte", "Refrigeración, fuelles, tipo camión"],
            ["vehicles / drivers", "vehiculos / conductores", "Satélite, fotos, solape horarios"],
            ["payrollEmployees", "empleados_nomina", "Periodicidad canónica vía API"],
            ["payrollRuns", "liquidaciones_nomina", "Automática, prima, terminación, intereses"],
            ["approvals", "solicitudes_autorizacion", "Ventana 10 minutos"],
            ["notifications", "notificaciones + preferencias", "Sonido configurable por usuario"],
        ],
    )
    add_p(doc, "Endpoints representativos del `PortalController`:", bold=True)
    add_bullets(
        doc,
        [
            "GET `/api/portal/bootstrap` — hidratación inicial post-login.",
            "POST `/api/portal/sync-key` — escritura incremental por clave.",
            "POST `/api/portal/approve-pending-user` — alta de empresas cliente.",
            "POST `/api/portal/transport-schedule-busy` — validación solapes de agenda.",
            "GET `/api/portal/candidates/:id/cv-download` — descarga segura de hojas de vida.",
        ],
    )

    add_h(doc, "8. Modelo de datos PostgreSQL", 1)
    add_bullets(
        doc,
        [
            f"Núcleo: `01_extensions` … `07_indices` — enums, empresas, usuarios, transporte, RRHH, sistema.",
            f"Evolución: {max(0, n_sql - 7)} migraciones incrementales (08–36+) — tarifas, RLS, nómina automática, B2B, SST.",
            "Índices de rendimiento: cobertura pública solicitudes, solape horarios viajes.",
            "Auditoría: `25_transporte_auditoria_eliminados` y soft-delete operativo en API.",
        ],
    )
    add_table(
        doc,
        ["Migración reciente", "Propósito"],
        [
            ["30_idx_solicitudes_transporte_cobertura_publica", "Consultas portal cliente"],
            ["33_empleados_examen_ocupacional_intravehicular", "Cumplimiento SST / intravehicular"],
            ["35_idx_viajes_horario_solape", "Calendario sin doble asignación"],
            ["36_drop_prospectos_b2b_volumen_mensual", "Alineación formulario B2B actual"],
        ],
    )

    doc.add_page_break()
    add_h(doc, "9. Atributos de calidad y pruebas", 1)
    add_table(
        doc,
        ["Atributo (ISO 25010)", "Implementación observable"],
        [
            ["Seguridad", "JWT, bcrypt, Turnstile, CSP meta, X-Frame-Options, CORS allowlist, Throttler"],
            ["Fiabilidad", "Pool PG con reconexión; bootstrap idempotente; auditoría eliminados"],
            ["Mantenibilidad", "Módulos Nest inyectables; dominios JS desacoplados; DDL versionado"],
            ["Compatibilidad", "Toggle ES/EN; formatos fecha/monetary regionalizados"],
            ["Portabilidad", "Docker Compose local; despliegue nube documentado"],
        ],
    )
    add_p(doc, "Automatización de regresión:", bold=True)
    add_bullets(
        doc,
        [
            "`qa/portal-regression-tests.mjs` — flujos críticos portal.",
            "`scripts/smoke-api.mjs` — salud API.",
            "`npm run verify` — build API/web, lint y QA en cadena.",
        ],
    )

    add_h(doc, "10. Roadmap, limitaciones y glosario", 1)
    add_table(
        doc,
        ["Tema pendiente / limitación", "Impacto", "Prioridad sugerida"],
        [
            ["OpenAPI/Swagger formal publicado", "Integraciones externas", "Media — pre-entrega final"],
            ["Unificación SPA vs `apps/web` Next", "Doble mantenimiento", "Baja — definir canónico"],
            ["Pruebas de carga (PEN) y threat modeling", "Riesgo producción", "Alta — pre go-live"],
            ["RLS Supabase en prod si aplica", "Seguridad multi-tenant", "Según arquitectura host"],
        ],
    )
    add_table(
        doc,
        ["Término", "Definición en Antares"],
        [
            ["AS-IS", "Estado actual del código en el repositorio, no diseño futuro."],
            ["Bootstrap", "Carga inicial servidor → caché local vía GET `/portal/bootstrap`."],
            ["KEYS.*", "Identificadores de colecciones en `app.js` / DomainRegistry."],
            ["Ventana 10 min", "Tiempo para aprobación manual antes de auto-aprobación."],
        ],
    )

    doc.add_page_break()
    add_h(doc, "Aprobación y trazabilidad", 1)
    add_table(
        doc,
        ["Rol", "Nombre", "Firma", "Fecha"],
        [
            ["Elaboró (arquitectura / desarrollo)", "JULIÁN DE JESÚS TOBÓN BOTERO", "", ""],
            ["Revisó (cliente / TI)", "TRANSPORTES ANTARES S.A.S.", "", ""],
            ["Aprobó (dirección)", "", "", ""],
        ],
    )
    add_mono(
        doc,
        f"Generado: {out.name}\nVersión {VERSION} · {_mes_es(FECHA_INFORME)}\n"
        f"Comando: python docs/generar_arquitectura_uiux.py\n"
        f"Diagramas: {fig_dir.resolve()}",
    )

    doc.save(out)
    print(f"DOCX -> {out.resolve()} ({out.stat().st_size // 1024} KB)")

    if args.no_pdf:
        return
    pdf_target = out.with_suffix(".pdf")
    slo = args.soffice.strip() or None

    def run_word():
        ok, m = export_pdf_via_word(out, pdf_target)
        return ok and pdf_target.is_file(), m

    def run_lo():
        ok, m = export_pdf_via_libreoffice(slo, out, out.parent)
        return ok and pdf_target.is_file(), m

    if args.libreoffice_first:
        ok, msg = run_lo()
        if not ok:
            ok, msg = run_word()
    else:
        ok, msg = run_word()
        if not ok:
            ok, msg = run_lo()
    if ok:
        print(f"PDF  -> {pdf_target.resolve()}")
    else:
        print(f"PDF no generado: {msg}")


if __name__ == "__main__":
    main()
