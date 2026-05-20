# -*- coding: utf-8 -*-
"""Genera Acta_Aprobacion_Hito_50_Desarrollo.docx — Hito 2, Cláusula Cuarta del contrato."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

_DOCS = Path(__file__).resolve().parent
_ROOT = Path(__file__).resolve().parents[2]
OUT = _DOCS / "Acta_Aprobacion_Hito_50_Desarrollo.docx"
FECHA_DOC = date(2026, 5, 16)


def _logo_path() -> Path | None:
    for p in (
        _ROOT / "imagenes empresa" / "Logo.png",
        _ROOT / "imagenes%20empresa" / "Logo.png",
    ):
        if p.is_file():
            return p
    return None


def add_h(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def add_p(doc: Document, text: str, bold: bool = False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val


def add_cover(doc: Document):
    for _ in range(3):
        doc.add_paragraph()
    logo = _logo_path()
    if logo:
        pg = doc.add_paragraph()
        pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pg.add_run().add_picture(str(logo), width=Inches(3.6))
    for text, pt, bold in [
        ("TRANSPORTES ANTARES S.A.S.", 20, True),
        ("Acta de aprobación escrita — Hito intermedio de desarrollo", 14, True),
        ("50 % de avance · Pago Hito 2 — Contrato de Desarrollo de Software", 12, True),
        ("Plataforma Web Integral para Empresa de Transporte de Carga", 11, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(pt)
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta.add_run(
        f"La Ceja, Antioquia — {FECHA_DOC.strftime('%d de %B de %Y').replace('May', 'mayo')}"
    )
    rm.italic = True
    rm.font.size = Pt(10)
    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = conf.add_run("DOCUMENTO CONTRACTUAL — Dos (2) ejemplares de igual tenor")
    rc.bold = True
    rc.font.size = Pt(10)
    rc.font.color.rgb = RGBColor(0xB7, 0x1C, 0x1C)
    doc.add_page_break()


def main():
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    add_cover(doc)

    add_h(doc, "1. Identificación del documento y marco contractual", 1)
    add_p(
        doc,
        "El presente instrumento constituye la APROBACIÓN ESCRITA de EL CONTRATANTE respecto del "
        "cumplimiento del cincuenta por ciento (50 %) del desarrollo de la Plataforma Web Integral "
        "objeto del Contrato de Desarrollo de Software celebrado entre TRANSPORTES ANTARES S.A.S. "
        "y JULIÁN DE JESÚS TOBÓN BOTERO, suscrito en La Ceja, Antioquia, el 24 de abril de 2026, "
        "cuyo valor total asciende a VEINTICINCO MILLONES DE PESOS COLOMBIANOS M/CTE "
        "($ 25.000.000 COP), de los cuales $ 20.000.000 COP corresponden al desarrollo y "
        "$ 5.000.000 COP al cierre del período de garantía.",
    )
    add_p(
        doc,
        "Esta acta habilita el reconocimiento y pago del HITO 2 — HITO INTERMEDIO previsto en la "
        "Cláusula Cuarta del contrato: treinta por ciento (30 %) del valor del desarrollo, "
        "equivalente a SEIS MILLONES DE PESOS COLOMBIANOS M/CTE ($ 6.000.000 COP), "
        "pagadero al completar el 50 % del desarrollo y con la aprobación escrita de "
        "EL CONTRATANTE que se formaliza mediante las firmas al final de este documento.",
    )
    add_table(
        doc,
        ["Referencia contractual", "Detalle"],
        [
            ["Contrato", "Contrato de Desarrollo de Software — Plataforma Web Integral"],
            ["Contratante", "TRANSPORTES ANTARES S.A.S. — NIT 900.830.324-2"],
            ["Representante legal", "Alfredo de Jesús Botero Castañeda — C.C. 15.383.391"],
            ["Contratista", "JULIÁN DE JESÚS TOBÓN BOTERO — C.C. 1.040.046.958"],
            ["Alcance integral", "Anexo A — cinco (5) módulos funcionales + stack y seguridad"],
            ["Plazo estimado", "Dos (2) meses calendario (inicio 25-abr-2026 · entrega est. 25-jun-2026)"],
            ["Hito que se aprueba", "Hito 2 — 30 % — $ 6.000.000 COP (50 % desarrollo + aprobación escrita)"],
            ["Repositorio de entrega técnica", "Monorepo «Antares Nuevo» (código fuente, BD/postgres, apps/api)"],
        ],
    )

    add_h(doc, "2. Partes", 1)
    add_p(doc, "EL CONTRATANTE:", bold=True)
    add_bullets(
        doc,
        [
            "Razón social: TRANSPORTES ANTARES S.A.S.",
            "NIT: 900.830.324-2",
            "Domicilio: La Ceja, Antioquia, Colombia",
            "Representante legal para efectos de esta acta: Alfredo de Jesús Botero Castañeda",
        ],
    )
    add_p(doc, "EL CONTRATISTA:", bold=True)
    add_bullets(
        doc,
        [
            "Nombre: JULIÁN DE JESÚS TOBÓN BOTERO",
            "C.C.: 1.040.046.958",
            "Calidad: Desarrollador de software independiente",
            "Correo de contacto contractual: judtobon3006@gmail.com",
            "Ciudad: La Ceja, Antioquia",
        ],
    )

    add_h(doc, "3. Metodología de medición del avance (50 %)", 1)
    add_p(
        doc,
        "Las partes acuerdan medir el porcentaje global de desarrollo conforme al cronograma por "
        "fases de la Cláusula Segunda del contrato y al desglose por módulos del Anexo A, "
        "utilizando los siguientes criterios objetivos:",
    )
    add_bullets(
        doc,
        [
            "Peso por fase contractual: Fase 1 (12,5 %), Fase 2 (37,5 %), Fase 3 (37,5 %), Fase 4 (12,5 %).",
            "Criterio de completitud: funcionalidad operativa en ambiente de desarrollo/pruebas, "
            "persistencia en PostgreSQL, integración API NestJS y evidencia en repositorio versionado.",
            "Informes quincenales de avance (Cláusula Tercera, numeral 3.1) como soporte de seguimiento.",
            "La aprobación no sustituye el Acta de Entrega y Recibo a Satisfacción del cierre (Hito 3).",
        ],
    )
    add_table(
        doc,
        ["Fase contractual", "Contenido (Cláusula Segunda)", "% fase", "Peso sobre total", "Aporte al 50 %"],
        [
            ["Fase 1", "Arquitectura, diseño UI/UX, modelo de datos", "100 %", "12,5 %", "12,5 %"],
            ["Fase 2", "Sitio público + portal clientes empresariales", "100 %", "37,5 %", "37,5 %"],
            ["Fase 3", "Panel administración + módulos RRHH (parcial)", "0 %*", "37,5 %", "0 %*"],
            ["Fase 4", "Integración, QA, ajustes y despliegue producción", "0 %", "12,5 %", "0 %"],
            ["TOTAL GLOBAL", "—", "—", "100 %", "50,0 %"],
        ],
    )
    add_p(
        doc,
        "* Nota: Existe avance técnico adelantado en componentes de la Fase 3 (estimado ~35–45 % de dicha fase) "
        "que las partes NO incorporan al cómputo del 50 % objeto de este hito, para mantener coherencia "
        "con el desglose contractual y reservar su validación formal al cierre de la Fase 3 y al Acta final.",
        bold=False,
    )

    doc.add_page_break()
    add_h(doc, "4. Detalle de entregables aprobados al 50 %", 1)

    add_h(doc, "4.1 Fase 1 — Arquitectura, diseño y modelo de datos", 2)
    add_bullets(
        doc,
        [
            "Documento de arquitectura técnica y sistema de diseño UX/UI (docs/Antares_Arquitectura_UI_UX.docx).",
            "Esquema relacional PostgreSQL versionado (carpeta BD/postgres/, migraciones 01–36+).",
            "API REST modular NestJS 10 (apps/api): auth JWT + refresh, portal bootstrap/sync, nómina, B2B, archivos.",
            "Arquitectura portal por capas (modules/portal, domain-bootstrap, persistencia híbrida API ↔ caché).",
            "Manual y runbook de despliegue (DEPLOYMENT_RUNBOOK.md, CONFIGURACION_SUPABASE_RENDER_POSTGRES.md).",
        ],
    )

    add_h(doc, "4.2 Fase 2 — Módulo 1: Sitio web corporativo público (Anexo A.1)", 2)
    add_table(
        doc,
        ["Requerimiento Anexo A", "Estado", "Evidencia / observación"],
        [
            ["Hero + llamada a contacto", "Implementado", "index.html — sección corporativa"],
            ["Quiénes somos (misión, visión, valores)", "Implementado", "Contenido editable vía portal admin"],
            ["Flota (tipos de vehículo, fichas)", "Implementado", "Presentación turbo, camión, tractocamión, etc."],
            ["Servicios y ventajas competitivas", "Implementado", "Cobertura nacional, flota propia, conductores"],
            ["Cobertura (rutas / ciudades)", "Implementado", "Listado territorial integrado"],
            ["Formulario contacto B2B", "Implementado", "API b2b-prospect + tabla prospectos_contacto_b2b"],
            ["Footer (redes, teléfono, dirección)", "Implementado", "Datos de empresa configurables"],
        ],
    )

    add_h(doc, "4.3 Fase 2 — Módulo 2: Portal clientes empresariales (Anexo A.2)", 2)
    add_table(
        doc,
        ["Requerimiento Anexo A", "Estado", "Evidencia / observación"],
        [
            ["Login, registro empresas, recuperación contraseña", "Implementado", "auth + aprobación admin de cuentas"],
            ["Sesiones JWT + refresh tokens", "Implementado", "apps/api/auth, columna refresh_token_hash"],
            ["Dashboard solicitudes activas", "Implementado", "Vista dashboard / solicitudes cliente"],
            ["Creación solicitud de viaje (origen, destino, carga, adjuntos)", "Implementado", "Módulo solicitudes + validaciones"],
            ["Estados Pendiente → Aprobada → En tránsito → Completada/Cancelada", "Implementado", "Flujo operativo en BD y UI"],
            ["Flujo aprobación 10 min + auto-aprobación", "Implementado", "Centro autorizaciones + temporizador"],
            ["Detalle viaje asignado (vehículo, conductor, ruta)", "Implementado", "Vista detalle y notificaciones"],
            ["Notificaciones al cliente en tiempo real", "Implementado", "Módulo notificaciones + preferencias usuario"],
        ],
    )

    add_h(doc, "4.4 Resumen por módulo contractual (ponderación referencial)", 2)
    add_table(
        doc,
        ["Módulo (Anexo A)", "Peso referencial", "% avance módulo", "Contribución"],
        [
            ["Módulo 1 — Sitio público", "20 %", "100 %", "20,0 %"],
            ["Módulo 2 — Portal clientes", "20 %", "100 %", "20,0 %"],
            ["Módulo 3 — Panel administración", "20 %", "0 %*", "0,0 %"],
            ["Módulo 4 — Nómina RRHH", "20 %", "0 %*", "0,0 %"],
            ["Módulo 5 — Contratación RRHH", "20 %", "0 %*", "0,0 %"],
            ["TOTAL (criterio fases 1–2)", "—", "—", "50,0 %"],
        ],
    )

    add_h(doc, "5. Declaración del contratista", 1)
    add_p(
        doc,
        "EL CONTRATISTA declara bajo la gravedad de juramento que, a la fecha de esta acta, ha ejecutado "
        "de manera diligente y conforme al Anexo A y a las buenas prácticas de ingeniería de software, "
        "el cincuenta por ciento (50 %) del desarrollo pactado, correspondiente a la finalización de "
        "las Fases 1 y 2 del cronograma contractual, con los entregables técnicos descritos en la "
        "sección 4; que el código se encuentra en repositorio privado accesible para revisión de "
        "EL CONTRATANTE; y que los informes de avance quincenales han reportado de forma consistente "
        "dicho porcentaje y los riesgos asociados al tramo restante.",
    )

    add_h(doc, "6. Aprobación escrita del contratante y autorización de pago", 1)
    add_p(
        doc,
        "Con fundamento en la revisión de los entregables relacionados, en los informes de avance "
        "presentados por EL CONTRATISTA y en el cumplimiento del umbral del cincuenta por ciento (50 %) "
        "del desarrollo conforme a la metodología acordada en la sección 3, EL CONTRATANTE:",
    )
    add_bullets(
        doc,
        [
            "APRUEBA de manera expresa y escrita el avance del 50 % del desarrollo de la Plataforma.",
            "AUTORIZA el pago del HITO 2 — HITO INTERMEDIO por valor de SEIS MILLONES DE PESOS "
            "COLOMBIANOS M/CTE ($ 6.000.000 COP), equivalente al treinta por ciento (30 %) del "
            "valor total del contrato, dentro de los plazos de la Cláusula Cuarta.",
            "Reconoce que el anticipo (Hito 1 — $ 6.000.000 COP) fue exigido antes del inicio del "
            "desarrollo conforme al contrato, sin perjuicio de que las partes verifiquen su "
            "constancia en los soportes contables respectivos.",
            "Manifiesta que la presente aprobación NO constituye recibo a satisfacción final ni "
            "inicio del período de garantía; ello ocurrirá con el Acta de Entrega del Hito 3.",
        ],
    )

    add_h(doc, "7. Alcance pendiente para el 50 % restante (compromiso de ejecución)", 1)
    add_p(
        doc,
        "Para transparencia del hito y continuidad del proyecto, las partes dejan constancia del "
        "trabajo programado que integrará el cincuenta por ciento (50 %) restante y los hitos 3 y 4:",
    )
    add_bullets(
        doc,
        [
            "Completar Fase 3: Panel de administración (bandeja solicitudes, flota, conductores, "
            "historial, reportes operativos) con validación formal de EL CONTRATANTE.",
            "Completar Módulo 4: Nómina — liquidación mensual, PDF desprendibles, exportación Excel, alertas.",
            "Completar Módulo 5: Contratación — vacantes, selección, entrevistas, contratos y gestión documental.",
            "Fase 4: Pruebas QA integrales, ajustes, despliegue producción y documentación OpenAPI/Swagger.",
            "Entrega final: Acta de Recibo a Satisfacción (Hito 3 — $ 8.000.000 COP) y garantía (Hito 4).",
        ],
    )

    add_h(doc, "8. Condiciones generales", 1)
    add_bullets(
        doc,
        [
            "Vigencia de confidencialidad y protección de datos (Cláusulas Novena y Ley 1581 de 2012).",
            "Propiedad intelectual: licencia de uso no exclusiva hasta pago total (Cláusula Sexta).",
            "Modificaciones de alcance solo por otrosí escrito (Cláusula Décima Cuarta).",
            "Solución de controversias: arreglo directo, conciliación y jurisdicción de La Ceja (Cláusula Décima Tercera).",
            "El presente documento se firma en dos (2) ejemplares del mismo tenor, uno para cada parte.",
        ],
    )

    add_h(doc, "9. Firmas", 1)
    add_p(doc, f"En La Ceja, Antioquia, a los {FECHA_DOC.day} días del mes de mayo de 2026.", bold=False)
    doc.add_paragraph()
    doc.add_paragraph()

    sig = doc.add_table(rows=2, cols=2)
    sig.style = "Table Grid"
    left = [
        "EL CONTRATANTE",
        "TRANSPORTES ANTARES S.A.S.",
        "NIT: 900.830.324-2",
        "",
        "_________________________________________",
        "Alfredo de Jesús Botero Castañeda",
        "Representante Legal",
        "C.C. 15.383.391",
        "",
        "Aprobación del 50 % del desarrollo y autorización de pago Hito 2 ($ 6.000.000 COP)",
    ]
    right = [
        "EL CONTRATISTA",
        "JULIÁN DE JESÚS TOBÓN BOTERO",
        "Desarrollador de Software independiente",
        "",
        "_________________________________________",
        "JULIÁN DE JESÚS TOBÓN BOTERO",
        "C.C. 1.040.046.958",
        "La Ceja, Antioquia",
        "",
        "Certificación de cumplimiento del 50 % del desarrollo",
    ]
    for i, line in enumerate(left):
        sig.rows[0].cells[0].text = "\n".join(left) if i == 0 else sig.rows[0].cells[0].text
    sig.rows[0].cells[0].text = "\n".join(left)
    sig.rows[0].cells[1].text = "\n".join(right)

    add_h(doc, "Anexo — Checklist de verificación para EL CONTRATANTE (opcional)", 1)
    add_table(
        doc,
        ["Ítem de verificación", "Sí", "No", "Observaciones"],
        [
            ["Sitio público navegable y formulario B2B operativo", "☐", "☐", ""],
            ["Registro/login cliente y creación de solicitud de viaje", "☐", "☐", ""],
            ["Flujo de aprobación con ventana de 10 minutos", "☐", "☐", ""],
            ["API y base de datos accesibles para demostración", "☐", "☐", ""],
            ["Documentación de arquitectura y despliegue entregada", "☐", "☐", ""],
        ],
    )

    doc.save(OUT)
    print(f"Generado: {OUT}")


if __name__ == "__main__":
    main()
