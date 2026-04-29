# -*- coding: utf-8 -*-
"""Genera Manual_Despliegue_Supabase_Cloudflare.docx en esta carpeta."""
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc: Document, text: str, bold: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main():
    out = Path(__file__).resolve().parent / "Manual_Despliegue_Supabase_Cloudflare.docx"
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(
        "Antares — Manual de despliegue: Supabase, Cloudflare, Vercel y Render"
    )
    r.bold = True
    r.font.size = Pt(16)

    doc.add_paragraph(
        "Documento generado para el repositorio Antares. Incluye reparto de responsabilidades "
        "entre proveedores, pasos de creación de la base de datos, documentos (Storage), "
        "variables de entorno y referencia a los scripts SQL (incluidos RLS)."
    )

    add_heading(doc, "1. Reparto recomendado", 1)
    add_para(
        doc,
        "Tabla resumen: qué va en cada capa cuando ya tienes Git, Vercel, Render y dominio en Cloudflare.",
    )
    doc.add_paragraph(
        "Postgres (esquema Antares), Auth opcional y Storage de documentos → Supabase.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "DNS, SSL, WAF, reglas de caché (sin cachear API) → Cloudflare.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Frontend → Vercel.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "API / lógica sensible (service_role, firmas, procesos largos) → Render.",
        style="List Bullet",
    )
    add_para(
        doc,
        "El rol service_role de Supabase ignora RLS; úsalo solo en el servidor (Render). "
        "Las políticas RLS protegen acceso con anon key o usuarios autenticados desde el cliente.",
    )

    add_heading(doc, "2. Creación del proyecto Supabase", 1)
    add_bullets(
        doc,
        [
            "Entrar en supabase.com → New project.",
            "Definir contraseña de base de datos (guardarla en gestor de secretos).",
            "Elegir región cercana a usuarios (p. ej. Sudamérica o US East según disponibilidad).",
            "Esperar estado Active.",
            "En Project Settings → API: anotar Project URL, anon key y service_role.",
            "En Project Settings → Database: connection string y pooling si la API usa Node en Render.",
        ],
    )

    add_heading(doc, "3. Aplicar el esquema SQL (tablas en español)", 1)
    add_para(
        doc,
        "Los identificadores de tablas y columnas del proyecto están definidos en español "
        "(sin tildes en nombres SQL). Los valores de ENUM tipo rol_usuario mantienen códigos "
        "alineados con la aplicación (p. ej. admin, client). Ejecutar en el SQL Editor en orden:",
    )
    add_bullets(
        doc,
        [
            "postgres/01_extensions.sql",
            "postgres/02_enums.sql",
            "postgres/03_nucleo_empresa_usuarios.sql",
            "postgres/04_transporte.sql",
            "postgres/05_rrhh.sql",
            "postgres/06_sistema.sql",
            "postgres/07_indices.sql",
            "Opcional: postgres/08_seed_tarifas_trayecto.sql",
            "postgres/09_rls_tablas.sql — Row Level Security en tablas públicas",
            "postgres/10_rls_storage_supabase.sql — políticas de Storage (tras crear buckets)",
        ],
    )

    add_heading(doc, "4. Row Level Security (RLS)", 1)
    add_para(
        doc,
        "El archivo 09_rls_tablas.sql habilita RLS y define políticas para los roles anon "
        "y authenticated de Supabase. Supone que el usuario del portal en public.usuarios "
        "tiene el mismo UUID que auth.users (auth.uid()) cuando uses Supabase Auth. Si solo "
        "conectas por API con service_role, el servidor no está limitado por estas políticas.",
    )
    add_para(
        doc,
        "Funciones auxiliares en el script: es_administrador_global(), es_rrhh(), "
        "id_empresa_usuario_actual(). Ajusta los roles según tu política de negocio.",
    )

    add_heading(doc, "5. Documentos y Storage", 1)
    add_bullets(
        doc,
        [
            "En Supabase → Storage: crear buckets privados (p. ej. documentos_contratos, documentos_adjuntos, documentos_rrhh).",
            "Ejecutar 10_rls_storage_supabase.sql después de crear los buckets con los mismos nombres (o adaptar políticas).",
            "Flujo recomendado: el navegador sube el archivo a tu API en Render; Render usa service_role y sube a Storage; guardas en Postgres la ruta y metadatos.",
            "Para descarga: URLs firmadas con caducidad o streaming desde la API.",
            "Cloudflare: no cachear rutas /api ni descargas privadas (reglas de caché).",
        ],
    )

    add_heading(doc, "6. Variables de entorno", 1)
    add_para(doc, "Vercel (solo públicas):", bold=True)
    add_bullets(
        doc,
        [
            "NEXT_PUBLIC_SUPABASE_URL (o prefijo que use tu framework)",
            "NEXT_PUBLIC_SUPABASE_ANON_KEY",
            "URL pública de la API si el front no llama a Supabase directamente",
        ],
    )
    add_para(doc, "Render (secretas):", bold=True)
    add_bullets(
        doc,
        [
            "SUPABASE_URL",
            "SUPABASE_SERVICE_ROLE_KEY",
            "Cadena de conexión directa a Postgres si tu API usa SQL además del cliente (usar URL con pooling según documentación Supabase).",
        ],
    )
    add_para(
        doc,
        "Nunca colocar service_role en variables NEXT_PUBLIC_* ni en código del navegador.",
    )

    add_heading(doc, "7. Cloudflare (dominio ya conectado)", 1)
    add_bullets(
        doc,
        [
            "Usar DNS/CNAME hacia Vercel y/o Render según subdominios (app, api).",
            "SSL/TLS: Full (strict) cuando el origen tenga certificado válido.",
            "Cache Rules: bypass para API y contenido autenticado.",
            "Rate limiting / WAF según necesidad.",
        ],
    )

    add_heading(doc, "8. Integración Vercel ↔ Supabase", 1)
    add_para(
        doc,
        "Opcional: integración Supabase en el marketplace de Vercel para inyectar variables; "
        "verificar que ninguna clave service_role quede expuesta al cliente.",
    )

    add_heading(doc, "9. Checklist final", 1)
    add_bullets(
        doc,
        [
            "Proyecto Supabase creado.",
            "Scripts 01–07 (+08 opcional) + 09 RLS + 10 Storage ejecutados sin error.",
            "Buckets Storage creados y políticas aplicadas.",
            "Variables en Render y Vercel configuradas.",
            "Prueba de login y lectura con usuario real (si usas Auth).",
        ],
    )

    doc.save(out)
    print(f"Escrito: {out}")


if __name__ == "__main__":
    main()
