# -*- coding: utf-8 -*-
"""
Convierte los manuales Markdown de docs/manuales-usuario/ a documentos Word (.docx)
con portada, estilos corporativos, tablas, listas e imágenes embebidas.

Uso:
  python docs/manuales-usuario/generar_manuales_word.py

Salida:
  docs/manuales-usuario/word/*.docx
  docs/manuales-usuario/word/Manual_Completo_Portal_Antares.docx
"""
from __future__ import annotations

import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
MANUALS_DIR = Path(__file__).resolve().parent
ASSETS_DIR = MANUALS_DIR / "assets"
OUT_DIR = MANUALS_DIR / "word"

C_BRAND = RGBColor(0x37, 0x7C, 0xC0)
C_BRAND_DARK = RGBColor(0x1E, 0x4A, 0x73)
C_TEXT = RGBColor(0x0B, 0x21, 0x38)
C_MUTED = RGBColor(0x3A, 0x5A, 0x78)
C_FILL_HDR = "377CC0"
C_FILL_ALT = "E8F4FC"
FECHA = date(2026, 7, 1)

MODULE_FILES = [
    ("00-introduccion.md", "Introducción general", "Acceso, navegación y elementos comunes del portal"),
    ("01-dashboard.md", "Dashboard", "Vista general operativa e indicadores del día"),
    ("02-solicitudes.md", "Mis solicitudes", "Radicación y seguimiento de solicitudes de transporte"),
    ("03-viajes.md", "Transporte · Viajes", "Asignación y gestión de viajes"),
    ("04-camiones.md", "Transporte · Camiones", "Catálogo de flota, combustible y taller"),
    ("05-conductores.md", "Transporte · Conductores", "Ficha operativa de conductores"),
    ("06-calendario.md", "Transporte · Calendario", "Agenda de viajes, entrevistas y novedades"),
    ("07-historial.md", "Historial y trazabilidad", "Auditoría de cambios del portal"),
    ("08-reporteria.md", "Centro de reportería", "Exportación de reportes y analítica operativa"),
    ("09-gestion-humana.md", "Gestión humana", "Nómina, colaboradores, ausencias y liquidaciones"),
    ("10-contratacion.md", "Contratación", "Vacantes, candidatos, entrevistas y contratos"),
    ("11-cumplimiento-laboral.md", "Cumplimiento laboral y SST", "Controles legales y de seguridad"),
    ("12-contacto-b2b.md", "Contacto web (B2B)", "Bandeja de prospectos comerciales"),
    ("13-usuarios-permisos.md", "Usuarios y permisos", "Administración de accesos al portal"),
    ("14-autorizaciones.md", "Centro de aprobaciones", "Bandejas de autorización centralizadas"),
    ("15-notificaciones.md", "Notificaciones", "Bandeja de avisos del sistema"),
    ("16-mi-perfil.md", "Mi perfil", "Datos personales y preferencias de cuenta"),
]


def _logo_path() -> Path | None:
    for p in (
        ROOT / "imagenes empresa" / "logo-recortado.png",
        ROOT / "imagenes empresa" / "Logo.png",
    ):
        if p.is_file():
            return p
    return None


def _mes_es(d: date) -> str:
    meses = (
        "enero", "febrero", "marzo", "abril", "mayo", "junio",
        "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
    )
    return f"{d.day} de {meses[d.month - 1]} de {d.year}"


def _set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _cell_write(cell, text: str, *, bold=False, size=10, color: RGBColor | None = None, align=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Calibri"
    if color:
        r.font.color.rgb = color


def _add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run2._r.append(instr)

    run3 = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run3._r.append(fld_sep)

    run4 = paragraph.add_run("1")
    run5 = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run5._r.append(fld_end)


def setup_document_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = C_TEXT

    for level, size, color in (
        (1, 16, C_BRAND_DARK),
        (2, 13, C_BRAND),
        (3, 12, C_BRAND_DARK),
    ):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def add_footer(doc: Document, label: str) -> None:
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.clear()
        r = p.add_run(f"Transportes Antares S.A.S — {label}  ·  Pág. ")
        r.font.size = Pt(9)
        r.font.name = "Calibri"
        r.font.color.rgb = C_MUTED
        _add_page_number_field(p)


def add_cover(doc: Document, title: str, subtitle: str, *, complete: bool = False) -> None:
    logo = _logo_path()
    for _ in range(2):
        doc.add_paragraph()

    if logo:
        pg = doc.add_paragraph()
        pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pg.add_run().add_picture(str(logo), width=Inches(3.8))

    band = doc.add_table(rows=1, cols=1)
    band.style = "Table Grid"
    cell = band.rows[0].cells[0]
    _set_cell_shading(cell, C_FILL_HDR)
    cell.text = ""

    main_title = "Manual de usuario — Portal empresarial" if complete else "Manual de usuario"
    for line, pt, bold, color in (
        ("TRANSPORTES ANTARES S.A.S", 20, True, RGBColor(0xFF, 0xFF, 0xFF)),
        (main_title, 14, True, RGBColor(0xE8, 0xF4, 0xFC)),
        (title, 13, True, RGBColor(0xCC, 0xE5, 0xF8)),
        (subtitle, 11, False, RGBColor(0xCC, 0xE5, 0xF8)),
    ):
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        r.bold = bold
        r.font.size = Pt(pt)
        r.font.name = "Calibri"
        if color:
            r.font.color.rgb = color

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta.add_run(f"Versión 1.0  ·  {_mes_es(FECHA)}  ·  Uso interno y capacitación")
    rm.font.size = Pt(10)
    rm.italic = True
    rm.font.color.rgb = C_MUTED
    rm.font.name = "Calibri"

    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    clean = _strip_md_links(text)
    h = doc.add_heading(clean, level=min(level, 3))
    for r in h.runs:
        r.font.name = "Calibri"
        r.font.color.rgb = C_BRAND_DARK if level == 1 else C_BRAND


def add_rich_paragraph(
    doc: Document,
    text: str,
    *,
    style: str | None = None,
    italic: bool = False,
    align=None,
    space_after: int = 6,
    bullet: bool = False,
) -> None:
    if bullet:
        p = doc.add_paragraph(style="List Bullet")
    elif style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    _append_inline_runs(p, text, default_italic=italic)


def _append_inline_runs(paragraph, text: str, *, default_italic: bool = False) -> None:
    pattern = re.compile(
        r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))"
    )
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.italic = default_italic
        token = m.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif token.startswith("[") and "](" in token:
            inner = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if inner:
                run = paragraph.add_run(inner.group(1))
                run.font.color.rgb = C_BRAND
        else:
            run = paragraph.add_run(token)
        run.font.name = run.font.name or "Calibri"
        run.font.size = run.font.size or Pt(11)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.italic = default_italic


def _strip_md_links(text: str) -> str:
    return re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)


def add_blockquote(doc: Document, text: str) -> None:
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    _set_cell_shading(cell, C_FILL_ALT)
    p = cell.paragraphs[0]
    _append_inline_runs(p, text.lstrip("> ").strip())
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    if not headers:
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for j, header in enumerate(headers):
        _cell_write(table.rows[0].cells[j], _strip_md_links(header), bold=True, size=10, color=C_BRAND_DARK)
        _set_cell_shading(table.rows[0].cells[j], C_FILL_ALT)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            _cell_write(table.rows[i].cells[j], _strip_md_links(val), size=10)
    doc.add_paragraph()


def add_image(doc: Document, src: Path, caption: str) -> None:
    if not src.is_file():
        add_rich_paragraph(doc, f"[Imagen no encontrada: {caption}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(src), width=Inches(6.2))
    if caption.strip():
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(caption.strip())
        cr.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = C_MUTED
        cr.font.name = "Calibri"
    doc.add_paragraph()


def parse_image_line(line: str) -> tuple[str, str] | None:
    m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
    if not m:
        return None
    return m.group(1), m.group(2)


def resolve_image_path(raw: str, md_path: Path) -> Path:
    rel = raw.strip()
    if rel.startswith("./"):
        rel = rel[2:]
    candidate = (md_path.parent / rel).resolve()
    if candidate.is_file():
        return candidate
    alt = (ASSETS_DIR / Path(rel).name).resolve()
    if alt.is_file():
        return alt
    parts = rel.replace("\\", "/").split("/")
    if len(parts) >= 2:
        alt2 = ASSETS_DIR / parts[-2] / parts[-1]
        if alt2.is_file():
            return alt2
    return candidate


def render_markdown(doc: Document, md_text: str, md_path: Path, *, skip_first_h1: bool = False) -> None:
    lines = md_text.splitlines()
    i = 0
    skipped_h1 = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("# ") and not stripped.startswith("## "):
            if skip_first_h1 and not skipped_h1:
                skipped_h1 = True
                i += 1
                continue
            add_heading(doc, stripped[2:].strip(), 1)
            i += 1
            continue

        if stripped.startswith("## "):
            add_heading(doc, stripped[3:].strip(), 2)
            i += 1
            continue

        if stripped.startswith("### "):
            add_heading(doc, stripped[4:].strip(), 3)
            i += 1
            continue

        img = parse_image_line(stripped)
        if img:
            caption, raw = img
            add_image(doc, resolve_image_path(raw, md_path), caption)
            i += 1
            continue

        if stripped.startswith("> "):
            add_blockquote(doc, stripped)
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            if len(table_lines) >= 2:
                headers = [c.strip() for c in table_lines[0].strip("|").split("|")]
                rows = []
                for tl in table_lines[2:]:
                    if re.match(r"^[\|\s:-]+$", tl):
                        continue
                    rows.append([c.strip() for c in tl.strip("|").split("|")])
                add_table(doc, headers, rows)
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            items = []
            while i < len(lines) and (lines[i].strip().startswith("- ") or lines[i].strip().startswith("* ")):
                items.append(lines[i].strip()[2:].strip())
                i += 1
            for item in items:
                add_rich_paragraph(doc, item, bullet=True, space_after=3)
            continue

        if re.match(r"^\d+\.\s", stripped):
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s", lines[i].strip()):
                items.append(re.sub(r"^\d+\.\s*", "", lines[i].strip()))
                i += 1
            for idx, item in enumerate(items, start=1):
                p = doc.add_paragraph(style="List Number")
                p.paragraph_format.space_after = Pt(3)
                _append_inline_runs(p, item)
            continue

        if stripped.startswith("[") and "⬅" in stripped or "➡" in stripped or "Volver al índice" in stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            _append_inline_runs(p, _strip_md_links(stripped))
            for r in p.runs:
                r.font.color.rgb = C_MUTED
                r.font.size = Pt(10)
            i += 1
            continue

        para_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (
                not nxt
                or nxt == "---"
                or nxt.startswith("#")
                or nxt.startswith("!")
                or nxt.startswith(">")
                or nxt.startswith("|")
                or nxt.startswith("- ")
                or nxt.startswith("* ")
                or re.match(r"^\d+\.\s", nxt)
                or (nxt.startswith("[") and ("⬅" in nxt or "➡" in nxt))
            ):
                break
            para_lines.append(nxt)
            i += 1
        add_rich_paragraph(doc, " ".join(para_lines))


def build_module_doc(md_file: str, title: str, subtitle: str) -> Document:
    md_path = MANUALS_DIR / md_file
    text = md_path.read_text(encoding="utf-8")
    doc = Document()
    setup_document_styles(doc)
    add_cover(doc, title, subtitle)
    render_markdown(doc, text, md_path, skip_first_h1=True)
    add_footer(doc, title)
    return doc


def build_complete_doc() -> Document:
    doc = Document()
    setup_document_styles(doc)
    add_cover(
        doc,
        "Manual completo del portal",
        "16 módulos · Guía de uso con capturas de pantalla",
        complete=True,
    )

    add_heading(doc, "Índice de módulos", 1)
    toc_rows = [[f"{i:02d}", title, subtitle] for i, (_, title, subtitle) in enumerate(MODULE_FILES, start=0)]
    add_table(doc, ["N.º", "Módulo", "Descripción"], toc_rows)
    add_rich_paragraph(
        doc,
        "Nota: en Microsoft Word puede insertar una tabla de contenido automática "
        "(Referencias → Tabla de contenido) usando los estilos Título 1–3 de este documento.",
        italic=True,
    )
    doc.add_page_break()

    for md_file, title, subtitle in MODULE_FILES:
        md_path = MANUALS_DIR / md_file
        text = md_path.read_text(encoding="utf-8")
        add_heading(doc, title, 1)
        add_rich_paragraph(doc, subtitle, italic=True, space_after=12)
        render_markdown(doc, text, md_path, skip_first_h1=True)
        doc.add_page_break()

    add_footer(doc, "Manual completo del portal")
    return doc


def build_index_doc() -> Document:
    readme = MANUALS_DIR / "README.md"
    doc = Document()
    setup_document_styles(doc)
    add_cover(doc, "Índice de manuales", "Catálogo de documentos de usuario del portal")
    render_markdown(doc, readme.read_text(encoding="utf-8"), readme, skip_first_h1=True)
    add_footer(doc, "Índice de manuales")
    return doc


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    generated: list[Path] = []

    for md_file, title, subtitle in MODULE_FILES:
        stem = Path(md_file).stem
        out = OUT_DIR / f"{stem}.docx"
        build_module_doc(md_file, title, subtitle).save(out)
        generated.append(out)
        print(f"  OK  {out.name}")

    complete = OUT_DIR / "Manual_Completo_Portal_Antares.docx"
    build_complete_doc().save(complete)
    generated.append(complete)
    print(f"  OK  {complete.name}")

    index = OUT_DIR / "Indice_Manuales.docx"
    build_index_doc().save(index)
    generated.append(index)
    print(f"  OK  {index.name}")

    print(f"\nGenerados {len(generated)} documentos Word en:\n  {OUT_DIR}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
